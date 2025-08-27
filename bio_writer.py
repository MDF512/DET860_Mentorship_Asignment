import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK


# === Config ===
excel_file = "F25 Mentor Bio Survey.xlsx"   # Path to your Excel survey
output_file = "Mentor_Bios.docx"

PHOTO_DIR = "./Question"  # folder where all photos are stored

# === Helpers ===
def clean(value):
    """Return None for NaN/empty; otherwise trimmed string."""
    if pd.isna(value):
        return None
    s = str(value).strip()
    return s if s else None

def find_photo(name: str) -> str | None:
    """Search PHOTO_DIR for a file containing the cadet's name."""
    if not name:
        return None
    #name_lower = name.lower().replace(" ", "_")  # normalize a bit
    for f in os.listdir(PHOTO_DIR):
        if name in f:
            return os.path.join(PHOTO_DIR, f)
    return None


def add_field(doc: Document, label: str, value):
    """Add a paragraph with bold label and value (N/A if missing)."""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    val = clean(value)
    p.add_run(val if val is not None else "N/A")


# === Load Excel ===
df = pd.read_excel(excel_file)

# === Build Document ===
doc = Document()
doc.add_heading("Air Force ROTC Mentor Bios", 0)

for i, row in df.iterrows():
    name = clean(row.get("What is your Name?")) or "Unknown"
    doc.add_heading(name, level=1)

    # Insert photo (if path exists)
    formal_name = clean(row.get("Name")) or "Unknown"
    photo_path = find_photo(formal_name)
    if photo_path and os.path.exists(photo_path):
        try:
            doc.add_picture(photo_path, width=Inches(2))
        except Exception as e:
            doc.add_paragraph(f"[Photo could not be added: {e}]")
    else:
        doc.add_paragraph("[No photo found]")

    # Bio Fields
    add_field(doc, "Major", row.get("Major"))
    add_field(doc, "Academic Year", row.get("Academic Year"))
    add_field(doc, "AFSC/SFSC", row.get("AFSC/SFSC"))
    add_field(doc, "Detachment Positions", row.get("Current and Past Detachment Positions"))
    add_field(doc, "Married", row.get("Are you Married?"))
    add_field(doc, "Hobbies", row.get("What are your Hobbies?"))
    add_field(doc, "Clubs", row.get("What clubs you are a part of?"))
    add_field(doc, "ROTC Strengths", row.get("What are your strengths in ROTC?"))
    add_field(doc, "Mentorship Style", row.get("What is your Mentorship style?"))
    add_field(doc, "Additional Info", row.get("Anything else you want to include about yourself:"))

    # Page break (unless last cadet)
    if i < len(df) - 1:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# === Save ===
doc.save(output_file)
print(f"Mentor bios exported to {output_file}")
