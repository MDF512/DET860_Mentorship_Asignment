import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import os

# === Config ===
excel_file = "mentor_survey.xlsx"   # <-- update to your file path
output_file = "Mentor_Bios.docx"

# === Helpers ===
def clean(value):
    """Return None for NaN/empty; otherwise trimmed string."""
    if pd.isna(value):
        return None
    s = str(value).strip()
    return s if s else None

def add_field(doc: Document, label: str, value):
    """Add a paragraph with bold label and value (N/A if missing)."""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    val = clean(value)
    p.add_run(val if val is not None else "N/A")

# === Load Excel ===
df = pd.read_excel(excel_file)

# === Build Document ===
doc = Document()
doc.add_heading("Air Force ROTC Mentor Bios", 0)

for _, row in df.iterrows():
    name = clean(row.get("""What is your Name?"""))
    doc.add_heading(name if name else "Unknown", level=1)

    name = clean(row.get("Name"))
    folder_path = "./Question"
    photo = [f for f in os.listdir(folder_path) if name.lower() in f.lower()]
    photo_path = folder_path + "/" + photo[0]
    doc.add_picture(photo_path, width=Inches(3))

    # Basic
    add_field(doc, "Major", row.get("Major"))
    add_field(doc, "Academic Year", row.get("Academic Year"))
    add_field(doc, "AFSC/SFSC", row.get("AFSC/SFSC"))

    # Detachment / Status
    add_field(doc, "Detachment Positions", row.get("Current and Past Detachment Positions"))
    add_field(doc, "Married", row.get("Are you Married?"))

    # Personal / Interests
    add_field(doc, "Hobbies", row.get("What are your Hobbies?"))
    # Handles NaN for Clubs gracefully:
    add_field(doc, "Clubs", row.get("What clubs you are a part of?"))

    # Strengths / Style
    add_field(doc, "ROTC Strengths", row.get("What are your strengths in ROTC?"))
    add_field(doc, "Mentorship Style", row.get("What is your Mentorship style?"))

    # Extra
    add_field(doc, "Additional Info", row.get("Anything else you want to include about yourself:"))

    # Spacer
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# === Save ===
doc.save(output_file)
print(f"Mentor bios exported to {output_file}")
